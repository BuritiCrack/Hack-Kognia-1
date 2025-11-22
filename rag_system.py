import os
from typing import List, Optional
from pathlib import Path
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader
import docx
import pickle


class DocumentProcessor:
    """Procesa y carga documentos PDF y DOCX."""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extrae texto de un archivo PDF."""
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extrae texto de un archivo DOCX."""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extrae texto de un archivo TXT."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @classmethod
    def process_document(cls, file_path: str, filename: str) -> str:
        """Procesa un documento y extrae su texto."""
        extension = Path(filename).suffix.lower()
        
        if extension == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return cls.extract_text_from_docx(file_path)
        elif extension == '.txt':
            return cls.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Formato de archivo no soportado: {extension}")


class RAGSystem:
    """Sistema RAG para consultas sobre documentos legales usando modelos locales."""
    
    def __init__(self, api_key: str, model_name: str = "huggingface"):
        self.api_key = api_key
        self.model_name = model_name
        
        # Usar embeddings locales gratuitos
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'}
        )
        
        self.vector_store: Optional[FAISS] = None
        self.documents_loaded = []
    
    def _expand_query(self, question: str) -> List[str]:
        """Expande la consulta con sinónimos y términos relacionados."""
        queries = [question]
        
        # Diccionario de expansiones para términos legales comunes
        expansions = {
            'requisitos': ['condiciones', 'requerimientos', 'exigencias', 'debe cumplir'],
            'obtener': ['conseguir', 'tramitar', 'solicitar', 'adquirir'],
            'licencia de conducción': ['licencia de conducir', 'pase de conducción', 'permiso de conducir'],
            'procedimiento': ['proceso', 'trámite', 'pasos', 'cómo hacer'],
            'obligaciones': ['deberes', 'responsabilidades', 'debe hacer'],
            'derechos': ['facultades', 'puede hacer', 'tiene derecho'],
            'sanciones': ['multas', 'penalidades', 'infracciones'],
            'plazo': ['término', 'tiempo', 'fecha límite', 'vencimiento'],
        }
        
        question_lower = question.lower()
        for term, synonyms in expansions.items():
            if term in question_lower:
                for synonym in synonyms:
                    expanded = question_lower.replace(term, synonym)
                    if expanded != question_lower:
                        queries.append(expanded)
        
        return queries[:3]  # Limitar a 3 variaciones
        
    def add_documents(self, texts: List[str], metadatas: List[dict]):
        """Añade documentos al sistema RAG."""
        # Crear splitter para dividir textos largos con mejor contexto
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,  # Chunks aún más grandes para capturar artículos completos
            chunk_overlap=400,  # Mayor overlap para no perder información entre chunks
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]  # Priorizar separadores naturales
        )
        
        # Crear documentos
        documents = []
        for text, metadata in zip(texts, metadatas):
            splits = text_splitter.split_text(text)
            for i, split in enumerate(splits):
                doc = Document(
                    page_content=split,
                    metadata={**metadata, "chunk": i}
                )
                documents.append(doc)
        
        self.documents_loaded.extend([m['filename'] for m in metadatas])
        
        # Crear o actualizar vector store
        if self.vector_store is None:
            self.vector_store = FAISS.from_documents(documents, self.embeddings)
        else:
            new_store = FAISS.from_documents(documents, self.embeddings)
            self.vector_store.merge_from(new_store)
        
    def query(self, question: str) -> dict:
        """Realiza una consulta al sistema RAG con búsqueda expandida y re-ranking."""
        if self.vector_store is None:
            raise ValueError("No hay documentos cargados en el sistema.")
        
        # Expandir la consulta con variaciones
        expanded_queries = self._expand_query(question)
        
        # Buscar con todas las variaciones y combinar resultados
        all_docs = {}
        for query in expanded_queries:
            docs_with_scores = self.vector_store.similarity_search_with_score(query, k=10)
            for doc, score in docs_with_scores:
                doc_key = doc.page_content[:100]  # Usar inicio como clave única
                if doc_key not in all_docs or all_docs[doc_key]['score'] > score:
                    all_docs[doc_key] = {
                        'doc': doc,
                        'score': score
                    }
        
        # Convertir a lista y ordenar por score
        docs_with_scores = [(item['doc'], item['score']) for item in all_docs.values()]
        docs_with_scores.sort(key=lambda x: x[1])  # Menor score = más similar
        
        # Re-ranking: priorizar chunks que contienen palabras clave de la pregunta
        keywords = self._extract_keywords(question)
        reranked_docs = self._rerank_by_keywords(docs_with_scores[:15], keywords)
        
        # Procesar fuentes con mejor re-ranking
        sources = []
        for doc, score, keyword_score in reranked_docs:
            sources.append({
                "content": doc.page_content,
                "filename": doc.metadata.get("filename", "Desconocido"),
                "chunk": doc.metadata.get("chunk", 0),
                "score": float(score),
                "keyword_score": keyword_score
            })
        
        # Filtrar resultados realmente relevantes
        relevant_sources = [s for s in sources if s['keyword_score'] >= 2]  # Al menos 2 keywords
        
        if not relevant_sources and sources:
            # Si el filtro es muy estricto, usar los mejores por score
            relevant_sources = sources[:3]
        
        # Crear una respuesta más estructurada y precisa
        if relevant_sources:
            # Analizar el tipo de pregunta para dar formato adecuado
            question_lower = question.lower()
            
            # Determinar el contexto de la pregunta
            is_definition = any(word in question_lower for word in ["qué es", "define", "significado", "concepto"])
            is_requirements = any(word in question_lower for word in ["requisitos", "condiciones", "debe cumplir", "exigencias"])
            is_list = any(word in question_lower for word in ["cuáles", "cuántos", "lista", "enumera"])
            is_how = any(word in question_lower for word in ["cómo", "de qué manera", "procedimiento", "proceso"])
            is_when = any(word in question_lower for word in ["cuándo", "plazo", "fecha", "tiempo"])
            is_amount = any(word in question_lower for word in ["cuánto", "valor", "precio", "monto", "costo"])
            
            # Construir respuesta según el tipo de pregunta
            answer_parts = []
            
            if is_definition:
                answer_parts.append("📖 **Definición Legal:**\n\n")
            elif is_requirements:
                answer_parts.append("📋 **Requisitos según la Ley:**\n\n")
            elif is_list:
                answer_parts.append("📋 **Información encontrada:**\n\n")
            elif is_how:
                answer_parts.append("⚙️ **Procedimiento Legal:**\n\n")
            elif is_when:
                answer_parts.append("📅 **Información sobre Plazos:**\n\n")
            elif is_amount:
                answer_parts.append("💰 **Valores/Montos:**\n\n")
            else:
                answer_parts.append("📄 **Según el documento legal:**\n\n")
            
            # Agrupar y mostrar los fragmentos más relevantes
            if is_requirements or is_list:
                # Para requisitos/listas, mostrar SOLO el fragmento más relevante
                best_source = relevant_sources[0]  # El primero es el más relevante (ya está ordenado)
                content = best_source['content'].strip()
                content_preview = self._create_smart_preview(content, keywords)
                
                relevance_text = f"{int(best_source['keyword_score'])} palabras clave"
                answer_parts.append(f"**Información más relevante** ({relevance_text}):\n\n{content_preview}\n\n")
                
                # Mencionar si hay más información disponible
                if len(relevant_sources) > 1:
                    answer_parts.append(f"_💡 Nota: Se encontraron {len(relevant_sources)} fragmentos relacionados. Mostrando el más relevante._\n")
            else:
                # Para otras preguntas, mostrar fragmentos individuales
                for i, source in enumerate(relevant_sources[:3], 1):
                    content = source['content'].strip()
                    content_preview = self._create_smart_preview(content, keywords)
                    
                    relevance_text = f"{int(source['keyword_score'])} palabras clave"
                    answer_parts.append(f"**Fragmento {i}** ({relevance_text}):\n\n{content_preview}\n\n---\n\n")
            
            # Agregar contexto adicional solo para preguntas que no sean requisitos/listas
            if not (is_requirements or is_list):
                if len(relevant_sources) > 5:
                    answer_parts.append(f"\n_💡 Se encontraron {len(relevant_sources)} fragmentos relevantes en total._\n")
                elif len(relevant_sources) > 3:
                    answer_parts.append(f"\n_Se encontraron {len(relevant_sources)} fragmentos relacionados._\n")
            
            answer = "".join(answer_parts)
        else:
            answer = "❌ **No se encontró información específica** para responder tu pregunta.\n\n"
            answer += "**Sugerencias:**\n"
            answer += "- Verifica que el documento contenga información sobre este tema\n"
            answer += "- Intenta usar términos diferentes (ej: 'requisitos' en vez de 'condiciones')\n"
            answer += "- Asegúrate de que el documento esté correctamente cargado\n"
        
        return {
            "answer": answer,
            "sources": relevant_sources if relevant_sources else sources[:3],
            "confidence": self._calculate_confidence(relevant_sources if relevant_sources else sources)
        }
    
    def _extract_keywords(self, question: str) -> List[str]:
        """Extrae palabras clave importantes de la pregunta."""
        # Palabras de relleno a ignorar
        stopwords = {'qué', 'cuál', 'cuáles', 'cómo', 'cuándo', 'dónde', 'quién', 'es', 'son', 'el', 'la', 'los', 'las',
                    'un', 'una', 'unos', 'unas', 'de', 'del', 'para', 'por', 'con', 'sin', 'sobre', 'según', 'debe',
                    'puede', 'tiene', 'hacer', 'ser', 'estar', 'en', 'a', 'y', 'o', 'u', 'e'}
        
        words = question.lower().split()
        keywords = [w.strip('¿?.,;:()') for w in words if w.lower() not in stopwords and len(w) > 3]
        
        # Agregar bigramas importantes
        bigrams = []
        for i in range(len(words) - 1):
            bigram = f"{words[i]} {words[i+1]}"
            if 'licencia' in bigram or 'requisito' in bigram or 'conducción' in bigram:
                bigrams.append(bigram)
        
        return keywords + bigrams
    
    def _rerank_by_keywords(self, docs_with_scores: List[tuple], keywords: List[str]) -> List[tuple]:
        """Re-rankea documentos basándose en la presencia de palabras clave."""
        scored_docs = []
        
        for doc, score in docs_with_scores:
            content_lower = doc.page_content.lower()
            
            # Contar keywords presentes
            keyword_score = sum(1 for kw in keywords if kw.lower() in content_lower)
            
            # Bonus por densidad de keywords
            if keyword_score > 0:
                keyword_density = keyword_score / len(content_lower.split()) * 1000
                keyword_score += keyword_density
            
            scored_docs.append((doc, score, keyword_score))
        
        # Ordenar por keyword_score (descendente) y luego por similarity score (ascendente)
        scored_docs.sort(key=lambda x: (-x[2], x[1]))
        
        return scored_docs
    
    def _create_smart_preview(self, content: str, keywords: List[str]) -> str:
        """Crea un preview inteligente mostrando el contexto completo y relevante."""
        content_lower = content.lower()
        
        # Buscar todas las posiciones de keywords
        keyword_positions = []
        for kw in keywords:
            pos = 0
            while True:
                pos = content_lower.find(kw.lower(), pos)
                if pos == -1:
                    break
                keyword_positions.append(pos)
                pos += 1
        
        if not keyword_positions:
            # Si no hay keywords, mostrar inicio completo sin truncar tanto
            return self._format_content(content[:1500])
        
        # Encontrar la posición más temprana de keywords
        best_position = min(keyword_positions)
        
        # Buscar el inicio del párrafo o sección más cercano
        start = content.rfind('\n\n', 0, best_position)
        if start == -1:
            # Si no hay doble salto, buscar salto simple
            start = content.rfind('\n', 0, best_position)
            if start == -1:
                start = max(0, best_position - 150)
        
        # Mostrar desde el inicio encontrado hasta el final del chunk
        preview = content[start:].strip()
        
        # Solo truncar si es realmente muy largo (más de 1800 chars)
        if len(preview) > 1800:
            # Buscar un punto de corte natural (punto final) después de 1500 chars
            cut_point = preview.find('. ', 1500)
            if cut_point != -1 and cut_point < 2000:
                preview = preview[:cut_point + 1]
            else:
                # Si no encuentra punto, buscar salto de línea
                cut_point = preview.find('\n', 1500)
                if cut_point != -1 and cut_point < 2000:
                    preview = preview[:cut_point]
                else:
                    preview = preview[:1800] + "..."
        
        return self._format_content(preview)
    
    def _format_content(self, text: str) -> str:
        """Formatea el contenido para mejor legibilidad."""
        import re
        
        # Primero limpiar espacios múltiples pero preservar saltos de línea importantes
        lines = text.split('\n')
        cleaned_lines = [' '.join(line.split()) for line in lines]
        text = ' '.join(cleaned_lines)
        
        # Agregar saltos de línea después de patrones específicos
        
        # Para artículos legales (Artículo 123.)
        text = re.sub(r'(Artículo\s+\d+[a-z]?\.)', r'\n\n\1', text, flags=re.IGNORECASE)
        
        # Para parágrafo (Parágrafo.)
        text = re.sub(r'(Parágrafo\s*\d*\.)', r'\n\n\1', text, flags=re.IGNORECASE)
        
        # Para enumeraciones con números seguidos de punto o paréntesis
        text = re.sub(r'\.\s+(\d+[.)])\s+', r'.\n\n\1 ', text)
        
        # Para enumeraciones con letras (a), b), c))
        text = re.sub(r'\.\s+([a-z][.)])\s+', r'.\n\n\1 ', text)
        
        # Para requisitos o items con guion o bullet
        text = re.sub(r'\.\s+([-•])\s+', r'.\n\n\1 ', text)
        
        # Para secciones (CAPITULO, TITULO, etc)
        text = re.sub(r'(CAPITULO\s+[IVXLCDM]+)', r'\n\n\1', text, flags=re.IGNORECASE)
        text = re.sub(r'(TITULO\s+[IVXLCDM]+)', r'\n\n\1', text, flags=re.IGNORECASE)
        
        # Limpiar espacios al inicio de líneas nuevas
        text = re.sub(r'\n\s+', r'\n', text)
        
        # Limpiar múltiples saltos de línea (máximo 2)
        text = re.sub(r'\n{3,}', r'\n\n', text)
        
        return text.strip()
    
    def _combine_related_fragments(self, sources: List[dict], keywords: List[str]) -> str:
        """Combina múltiples fragmentos relacionados en una respuesta cohesiva."""
        combined_parts = []
        seen_content = set()
        
        for i, source in enumerate(sources, 1):
            content = source['content'].strip()
            
            # Evitar duplicados (comparar primeros 200 chars)
            content_signature = content[:200]
            if content_signature in seen_content:
                continue
            seen_content.add(content_signature)
            
            # Crear preview inteligente
            preview = self._create_smart_preview(content, keywords)
            
            # Detectar si contiene lista/enumeración
            has_enumeration = any(pattern in preview for pattern in ['1.', '2.', 'a)', 'b)', '•', '-'])
            
            relevance_text = f"{int(source['keyword_score'])} palabras clave"
            
            if has_enumeration:
                # Formato especial para listas
                combined_parts.append(f"**Sección {i}** ({relevance_text}):\n\n{preview}\n\n---\n")
            else:
                # Formato normal
                combined_parts.append(f"**Fragmento {i}** ({relevance_text}):\n\n{preview}\n\n---\n")
        
        return "\n".join(combined_parts)
    
    def _calculate_confidence(self, sources: List[dict]) -> str:
        """Calcula el nivel de confianza basado en la cantidad y calidad de fuentes."""
        if not sources:
            return "Baja"
        
        # Calcular confianza basada en scores (FAISS usa distancia L2, menor es mejor)
        avg_score = sum(s.get('score', 100) for s in sources[:3]) / min(len(sources), 3)
        
        # Para embeddings de 384 dimensiones, scores típicos:
        # < 0.5 = Muy similar (Alta confianza)
        # 0.5-0.8 = Similar (Media confianza)
        # > 0.8 = Poco similar (Baja confianza)
        if avg_score < 0.5 and len(sources) >= 3:
            return "Alta"
        elif avg_score < 0.8 and len(sources) >= 2:
            return "Media"
        else:
            return "Baja"
    
    def save_vector_store(self, path: str):
        """Guarda el vector store en disco."""
        if self.vector_store:
            self.vector_store.save_local(path)
    
    def load_vector_store(self, path: str):
        """Carga el vector store desde disco."""
        self.vector_store = FAISS.load_local(
            path,
            self.embeddings,
            allow_dangerous_deserialization=True
        )
    
    def reset(self):
        """Reinicia el sistema RAG."""
        self.vector_store = None
        self.documents_loaded = []
